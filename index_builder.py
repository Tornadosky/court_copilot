"""
Document indexing module for building FAISS vector store from legal documents.
Processes PDFs, Word docs, and text files into searchable chunks.
"""

import os
import re
import json
from pathlib import Path
from typing import List, Dict, Tuple, Optional
import asyncio
from concurrent.futures import ThreadPoolExecutor

# Document processing imports
import fitz  # PyMuPDF
from docx import Document
from pdfminer.high_level import extract_text as pdf_extract_text

# Vector store and embedding imports
import numpy as np
import faiss
from openai import AsyncOpenAI
import tiktoken

from rich.console import Console
from rich.progress import Progress, TextColumn, BarColumn, TaskProgressColumn, TimeRemainingColumn

console = Console()

class DocumentIndexer:
    """
    Document indexer that processes legal documents and builds a FAISS vector store.
    Supports PDF, Word, and text documents with intelligent chunking.
    """
    
    def __init__(self, api_key: str, embedding_model: str = "text-embedding-3-small",
                 chunk_size: int = 300, chunk_overlap: int = 50):
        """
        Initialize the document indexer.
        
        Args:
            api_key: OpenAI API key for embeddings
            embedding_model: OpenAI embedding model to use
            chunk_size: Target chunk size in tokens
            chunk_overlap: Overlap between chunks in tokens
        """
        self.client = AsyncOpenAI(api_key=api_key)
        self.embedding_model = embedding_model
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Initialize tokenizer for chunk size estimation
        self.tokenizer = tiktoken.encoding_for_model("gpt-4")
        
        # Document storage
        self.documents = []  # List of processed documents
        self.chunks = []     # List of text chunks
        self.embeddings = None  # Numpy array of embeddings
        self.index = None    # FAISS index
        
        console.print(f"[green]Document indexer initialized:[/green] {embedding_model}, {chunk_size} tokens/chunk")
    
    def extract_text_from_pdf(self, file_path: str) -> Tuple[str, Dict]:
        """
        Extract text from PDF file using PyMuPDF (fast) with fallback to pdfminer.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            # Try PyMuPDF first (faster)
            doc = fitz.open(file_path)
            text_parts = []
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text()
                if text.strip():  # Only add non-empty pages
                    text_parts.append(f"[Page {page_num + 1}]\n{text}")
            
            doc.close()
            
            full_text = "\n\n".join(text_parts)
            metadata = {
                "pages": len(doc),
                "extraction_method": "PyMuPDF"
            }
            
            return full_text, metadata
            
        except Exception as e:
            console.print(f"[yellow]PyMuPDF failed for {file_path}, trying pdfminer: {e}[/yellow]")
            
            try:
                # Fallback to pdfminer
                text = pdf_extract_text(file_path)
                metadata = {
                    "pages": "unknown",
                    "extraction_method": "pdfminer"
                }
                return text, metadata
                
            except Exception as e2:
                console.print(f"[red]Both PDF extraction methods failed for {file_path}: {e2}[/red]")
                return "", {"error": str(e2)}
    
    def extract_text_from_docx(self, file_path: str) -> Tuple[str, Dict]:
        """
        Extract text from Word document.
        
        Args:
            file_path: Path to Word document
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            doc = Document(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            full_text = "\n\n".join(text_parts)
            metadata = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables)
            }
            
            return full_text, metadata
            
        except Exception as e:
            console.print(f"[red]Failed to extract text from {file_path}: {e}[/red]")
            return "", {"error": str(e)}
    
    def extract_text_from_txt(self, file_path: str) -> Tuple[str, Dict]:
        """
        Extract text from plain text file.
        
        Args:
            file_path: Path to text file
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            metadata = {
                "lines": len(text.split('\n')),
                "encoding": "utf-8"
            }
            
            return text, metadata
            
        except UnicodeDecodeError:
            try:
                # Try different encodings
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                
                metadata = {
                    "lines": len(text.split('\n')),
                    "encoding": "latin-1"
                }
                
                return text, metadata
                
            except Exception as e:
                console.print(f"[red]Failed to read text file {file_path}: {e}[/red]")
                return "", {"error": str(e)}
    
    def process_document(self, file_path: str) -> Optional[Dict]:
        """
        Process a single document and extract text.
        
        Args:
            file_path: Path to document file
            
        Returns:
            Document dictionary or None if processing failed
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            console.print(f"[red]File not found: {file_path}[/red]")
            return None
        
        # Determine file type and extract text
        file_extension = file_path.suffix.lower()
        
        if file_extension == '.pdf':
            text, metadata = self.extract_text_from_pdf(str(file_path))
        elif file_extension in ['.docx', '.doc']:
            text, metadata = self.extract_text_from_docx(str(file_path))
        elif file_extension in ['.txt', '.md']:
            text, metadata = self.extract_text_from_txt(str(file_path))
        else:
            console.print(f"[yellow]Unsupported file type: {file_extension}[/yellow]")
            return None
        
        if not text.strip():
            console.print(f"[yellow]No text extracted from {file_path}[/yellow]")
            return None
        
        # Create document record
        document = {
            "file_path": str(file_path),
            "file_name": file_path.name,
            "file_type": file_extension,
            "text": text,
            "metadata": metadata,
            "token_count": len(self.tokenizer.encode(text)),
            "char_count": len(text)
        }
        
        console.print(f"[green]✅ Processed:[/green] {file_path.name} ({document['token_count']} tokens)")
        
        return document
    
    def chunk_text(self, text: str, document_id: int) -> List[Dict]:
        """
        Split text into overlapping chunks for better retrieval.
        
        Args:
            text: Full document text
            document_id: ID of the source document
            
        Returns:
            List of chunk dictionaries
        """
        # Split text into sentences (rough approximation)
        sentences = re.split(r'[.!?]\s+', text)
        
        chunks = []
        current_chunk = []
        current_tokens = 0
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            
            # Estimate tokens in sentence
            sentence_tokens = len(self.tokenizer.encode(sentence))
            
            # Check if adding this sentence would exceed chunk size
            if current_tokens + sentence_tokens > self.chunk_size and current_chunk:
                # Create chunk from current sentences
                chunk_text = ". ".join(current_chunk) + "."
                chunks.append({
                    "text": chunk_text,
                    "document_id": document_id,
                    "token_count": current_tokens,
                    "chunk_id": len(chunks)
                })
                
                # Start new chunk with overlap
                overlap_sentences = current_chunk[-max(1, self.chunk_overlap // 20):]  # Rough overlap
                current_chunk = overlap_sentences + [sentence]
                current_tokens = sum(len(self.tokenizer.encode(s)) for s in current_chunk)
            else:
                current_chunk.append(sentence)
                current_tokens += sentence_tokens
        
        # Add final chunk if it has content
        if current_chunk:
            chunk_text = ". ".join(current_chunk) + "."
            chunks.append({
                "text": chunk_text,
                "document_id": document_id,
                "token_count": current_tokens,
                "chunk_id": len(chunks)
            })
        
        return chunks
    
    async def get_embeddings(self, texts: List[str]) -> np.ndarray:
        """
        Get embeddings for a list of texts using OpenAI API.
        
        Args:
            texts: List of text strings to embed
            
        Returns:
            Numpy array of embeddings
        """
        if not texts:
            return np.array([])
        
        try:
            # Split into batches to avoid API limits
            batch_size = 100
            all_embeddings = []
            
            for i in range(0, len(texts), batch_size):
                batch = texts[i:i + batch_size]
                
                response = await self.client.embeddings.create(
                    model=self.embedding_model,
                    input=batch
                )
                
                batch_embeddings = [item.embedding for item in response.data]
                all_embeddings.extend(batch_embeddings)
            
            return np.array(all_embeddings)
            
        except Exception as e:
            console.print(f"[red]Failed to get embeddings: {e}[/red]")
            raise
    
    def build_faiss_index(self, embeddings: np.ndarray) -> faiss.Index:
        """
        Build FAISS index from embeddings.
        
        Args:
            embeddings: Numpy array of embeddings
            
        Returns:
            FAISS index
        """
        if embeddings.size == 0:
            raise ValueError("No embeddings to index")
        
        # Create FAISS index (using flat L2 distance for simplicity)
        dimension = embeddings.shape[1]
        index = faiss.IndexFlatL2(dimension)
        
        # Add embeddings to index
        index.add(embeddings.astype(np.float32))
        
        console.print(f"[green]FAISS index built:[/green] {index.ntotal} vectors, {dimension} dimensions")
        
        return index
    
    async def index_documents(self, docs_directory: str) -> None:
        """
        Index all documents in a directory.
        
        Args:
            docs_directory: Path to directory containing documents
        """
        docs_path = Path(docs_directory)
        
        if not docs_path.exists():
            console.print(f"[red]Documents directory not found: {docs_directory}[/red]")
            return
        
        # Find all supported document files
        supported_extensions = ['.pdf', '.docx', '.doc', '.txt', '.md']
        doc_files = []
        
        for ext in supported_extensions:
            doc_files.extend(docs_path.glob(f"**/*{ext}"))
        
        if not doc_files:
            console.print(f"[yellow]No supported documents found in {docs_directory}[/yellow]")
            return
        
        console.print(f"[cyan]Found {len(doc_files)} documents to process[/cyan]")
        
        # Process documents in parallel
        with Progress(
            TextColumn("[bold blue]Processing documents..."),
            BarColumn(),
            TaskProgressColumn(),
            TimeRemainingColumn(),
            console=console
        ) as progress:
            
            task = progress.add_task("Processing", total=len(doc_files))
            
            # Use thread pool for I/O-bound document processing
            with ThreadPoolExecutor(max_workers=4) as executor:
                futures = [
                    executor.submit(self.process_document, str(doc_file))
                    for doc_file in doc_files
                ]
                
                for future in futures:
                    document = future.result()
                    if document:
                        self.documents.append(document)
                    progress.advance(task)
        
        if not self.documents:
            console.print("[red]No documents were successfully processed[/red]")
            return
        
        # Create chunks from all documents
        console.print("[cyan]Creating text chunks...[/cyan]")
        
        for doc_id, document in enumerate(self.documents):
            doc_chunks = self.chunk_text(document["text"], doc_id)
            self.chunks.extend(doc_chunks)
        
        console.print(f"[green]Created {len(self.chunks)} text chunks[/green]")
        
        # Generate embeddings
        console.print("[cyan]Generating embeddings...[/cyan]")
        
        chunk_texts = [chunk["text"] for chunk in self.chunks]
        self.embeddings = await self.get_embeddings(chunk_texts)
        
        # Build FAISS index
        console.print("[cyan]Building FAISS index...[/cyan]")
        self.index = self.build_faiss_index(self.embeddings)
        
        console.print(f"[bold green]✅ Indexing complete![/bold green]")
        console.print(f"  Documents: {len(self.documents)}")
        console.print(f"  Chunks: {len(self.chunks)}")
        console.print(f"  Embeddings: {self.embeddings.shape}")
    
    def save_index(self, index_directory: str) -> None:
        """
        Save the index and metadata to disk.
        
        Args:
            index_directory: Directory to save index files
        """
        index_path = Path(index_directory)
        index_path.mkdir(exist_ok=True)
        
        if self.index is None:
            console.print("[red]No index to save[/red]")
            return
        
        # Save FAISS index
        faiss.write_index(self.index, str(index_path / "faiss.index"))
        
        # Save metadata
        metadata = {
            "documents": self.documents,
            "chunks": self.chunks,
            "embedding_model": self.embedding_model,
            "chunk_size": self.chunk_size,
            "chunk_overlap": self.chunk_overlap
        }
        
        with open(index_path / "metadata.json", 'w') as f:
            json.dump(metadata, f, indent=2)
        
        console.print(f"[green]Index saved to {index_directory}[/green]")
    
    def load_index(self, index_directory: str) -> bool:
        """
        Load a previously saved index.
        
        Args:
            index_directory: Directory containing saved index
            
        Returns:
            True if loaded successfully, False otherwise
        """
        index_path = Path(index_directory)
        
        faiss_file = index_path / "faiss.index"
        metadata_file = index_path / "metadata.json"
        
        if not (faiss_file.exists() and metadata_file.exists()):
            console.print(f"[red]Index files not found in {index_directory}[/red]")
            return False
        
        try:
            # Load FAISS index
            self.index = faiss.read_index(str(faiss_file))
            
            # Load metadata
            with open(metadata_file, 'r') as f:
                metadata = json.load(f)
            
            self.documents = metadata["documents"]
            self.chunks = metadata["chunks"]
            self.embedding_model = metadata["embedding_model"]
            self.chunk_size = metadata["chunk_size"]
            self.chunk_overlap = metadata["chunk_overlap"]
            
            console.print(f"[green]Index loaded from {index_directory}[/green]")
            console.print(f"  Documents: {len(self.documents)}")
            console.print(f"  Chunks: {len(self.chunks)}")
            console.print(f"  Index size: {self.index.ntotal}")
            
            return True
            
        except Exception as e:
            console.print(f"[red]Failed to load index: {e}[/red]")
            return False


async def main():
    """Main function for building document index."""
    import argparse
    import os
    
    parser = argparse.ArgumentParser(description="Build document index for Courtroom AI Assistant")
    parser.add_argument("--docs", default="./docs", help="Directory containing documents to index")
    parser.add_argument("--index", default="./index", help="Directory to save index files")
    parser.add_argument("--chunk-size", type=int, default=300, help="Chunk size in tokens")
    parser.add_argument("--chunk-overlap", type=int, default=50, help="Chunk overlap in tokens")
    
    args = parser.parse_args()
    
    # Get API key
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        console.print("[red]OPENAI_API_KEY environment variable not set[/red]")
        return
    
    # Create indexer and process documents
    indexer = DocumentIndexer(
        api_key=api_key,
        chunk_size=args.chunk_size,
        chunk_overlap=args.chunk_overlap
    )
    
    await indexer.index_documents(args.docs)
    
    if indexer.index is not None:
        indexer.save_index(args.index)
    else:
        console.print("[red]Failed to create index[/red]")


if __name__ == "__main__":
    asyncio.run(main()) 